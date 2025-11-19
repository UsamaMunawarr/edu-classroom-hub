# ==========================================
#  Backend + Login
# ==========================================

import streamlit as st
import pandas as pd
import json, time, random, string, re, os, difflib
from datetime import datetime, timedelta
from io import BytesIO
from docx import Document
from docx.shared import Pt
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet
from PyPDF2 import PdfReader
import sqlite3
import os
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from streamlit_autorefresh import st_autorefresh

# Optional AI / Embeddings
try:
    import google.generativeai as genai
    from sentence_transformers import SentenceTransformer, util
    HAVE_GEN = True
except Exception:
    HAVE_GEN = False

# -------------------------
# CONFIGURATION
# -------------------------
st.set_page_config(page_title="🎓 EduGenius Hub", layout="wide")

TEACHER_PASS = "teach123"
STUDENT_PASS = "stud123"

# -------------------------
# SQLITE DATABASE
# -------------------------
conn = sqlite3.connect("results.db", check_same_thread=False)
c = conn.cursor()
c.execute("""
CREATE TABLE IF NOT EXISTS results (
    id INTEGER PRIMARY KEY AUTOINCREMENT,
    name TEXT, email TEXT,
    score REAL, total REAL,
    breakdown TEXT, timestamp TEXT
)
""")
conn.commit()

# -------------------------
# GEMINI & EMBEDDINGS
# -------------------------
from dotenv import load_dotenv
load_dotenv()

# API_KEY = os.getenv("API_KEY")
API_KEY = st.secrets.get("API_KEY")
# Gemini setup (optional AI)
if HAVE_GEN:
    try:
        genai.configure(api_key=API_KEY)
        try:
            gen_model = genai.GenerativeModel("models/gemini-2.5-pro")
        except Exception:
            gen_model = genai.GenerativeModel("models/gemini-2.5-flash")
    except Exception:
        gen_model = None
else:
    gen_model = None

# Don't load SentenceTransformer here! We'll lazy-load it inside grading function
embed_model = None

# -------------------------
# UTILITY FUNCTIONS
# -------------------------
def extract_text_from_pdf(uploaded_file) -> str:
    reader = PdfReader(uploaded_file)
    text = ""
    for p in reader.pages:
        text += p.extract_text() or ""
    return text

def generate_question_paper_from_text(text, num_mcq=5, num_short=5, num_long=2, style=""):
    if gen_model is None:
        mcqs = [
            {"id": i+1, "question": f"Demo MCQ #{i+1}",
             "options": ["A", "B", "C", "D"], "answer_letter": "a"}
            for i in range(num_mcq)
        ]
        shorts = [{"id": i+1, "question": f"Demo Short #{i+1}", "reference_answer": "Sample"} for i in range(num_short)]
        longs = [{"id": i+1, "question": f"Demo Long #{i+1}", "reference_answer": "Sample"} for i in range(num_long)]
        return {"mcqs": mcqs, "shorts": shorts, "longs": longs}
    prompt = f"""
You are an exam designer. Create JSON with keys mcqs, shorts, longs.
Each mcq: {{id,question,options:[a,b,c,d],answer_letter}}
Each short/long: {{id,question,reference_answer}}
Generate exactly {num_mcq} mcqs, {num_short} shorts, {num_long} longs.
Text: {text[:4000]}  Style: {style}
Return ONLY valid JSON.
"""
    resp = gen_model.generate_content(prompt)
    txt = resp.text.strip()
    start, end = txt.find('{'), txt.rfind('}') + 1
    try:
        data = json.loads(txt[start:end])
    except Exception:
        return generate_question_paper_from_text(text, num_mcq, num_short, num_long, style)
    for k in ["mcqs", "shorts", "longs"]:
        for i, q in enumerate(data.get(k, []), start=1):
            q.setdefault("id", i)
    return data


def score_short_answer(student_answer, reference, max_marks=5):
    global embed_model

    # Lazy-load embeddings only if needed
    if embed_model is None:
        try:
            from sentence_transformers import SentenceTransformer, util
            embed_model = SentenceTransformer("all-MiniLM-L6-v2")
        except Exception:
            embed_model = None

    # If embeddings unavailable, fallback
    if embed_model is None:
        sim = 1.0 if reference.lower().strip() in student_answer.lower() else 0.0
        marks = max_marks if sim >= 0.75 else max_marks * 0.5 if sim >= 0.4 else 0.0
        return marks, sim

    # Use embeddings to calculate similarity
    emb1 = embed_model.encode(student_answer, convert_to_tensor=True)
    emb2 = embed_model.encode(reference, convert_to_tensor=True)
    sim = util.pytorch_cos_sim(emb1, emb2).item()

    if sim >= 0.75: marks = max_marks
    elif sim >= 0.6: marks = max_marks * 0.8
    elif sim >= 0.45: marks = max_marks * 0.5
    elif sim >= 0.3: marks = max_marks * 0.25
    else: marks = 0.0

    return marks, sim


def grade_long_answer_fallback(question, reference_answer, student_answer, max_marks=10):
    global gen_model, embed_model

    try:
        if gen_model is None:
            raise Exception("Gemini not configured")
        prompt = f"""
You are a strict but fair grader. Return JSON: {{score: number (0-{max_marks}), feedback: string}}.
Question: {question}
Reference answer: {reference_answer}
Student answer: {student_answer}
"""
        resp = gen_model.generate_content(prompt).text.strip()
        parsed = json.loads(resp)
        score = float(parsed.get("score", 0))
        feedback = parsed.get("feedback", "")
        return max(0.0, min(score, max_marks)), feedback
    except Exception:
        # Lazy-load embeddings if available
        if embed_model is None:
            try:
                from sentence_transformers import SentenceTransformer, util
                embed_model = SentenceTransformer("all-MiniLM-L6-v2")
            except Exception:
                embed_model = None

        # Fallback using embeddings
        if embed_model is None:
            try:
                vectorizer = TfidfVectorizer().fit([reference_answer, student_answer])
                vectors = vectorizer.transform([reference_answer, student_answer])
                sim = cosine_similarity(vectors[0:1], vectors[1:2])[0][0]
                return round(sim * max_marks, 2), f"Auto-graded similarity {sim*100:.1f}% (fallback)"
            except Exception as e2:
                return 0.0, f"Auto-grade failed: {str(e2)}"
        else:
            emb1 = embed_model.encode(student_answer, convert_to_tensor=True)
            emb2 = embed_model.encode(reference_answer, convert_to_tensor=True)
            sim = util.pytorch_cos_sim(emb1, emb2).item()
            marks = min(max(sim * max_marks, 0.0), max_marks)
            return marks, f"Auto-graded similarity {sim*100:.1f}%"





def grade_long_answer(student_answer, reference, max_marks=10):
    global embed_model
    from sentence_transformers import SentenceTransformer, util
    if embed_model is None:
        embed_model = SentenceTransformer("all-MiniLM-L6-v2")
    emb1 = embed_model.encode(student_answer, convert_to_tensor=True)
    emb2 = embed_model.encode(reference, convert_to_tensor=True)
    sim = util.pytorch_cos_sim(emb1, emb2).item()
    marks = min(max(sim * max_marks, 0), max_marks)
    return score_short_answer(student_answer, reference, max_marks)

def make_pdf_report(data, student_answers, marks_config):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf)
    styles = getSampleStyleSheet()
    story = [Paragraph("Own Exam Report", styles["Title"]), Spacer(1,12)]

    # MCQs
    story.append(Paragraph("Section A - MCQs", styles["Heading2"]))
    for q in data.get("mcqs", []):
        story.append(Paragraph(f"Q{q['id']}: {q['question']}", styles["Normal"]))
        for opt in q.get("options", []):
            story.append(Paragraph(f"- {opt}", styles["Normal"]))
        student_ans = student_answers.get("mcq", {}).get(str(q['id']), "")
        correct_ans = q.get("options",[])[ord(q.get("answer_letter","a"))-97]
        mark = marks_config.get("mcq",1) if student_ans==correct_ans else 0
        story.append(Paragraph(f"Reference Answer: {correct_ans}", styles["Normal"]))
        story.append(Paragraph(f"Student Answer: {student_ans}", styles["Normal"]))
        story.append(Paragraph(f"Marks: {mark}/{marks_config.get('mcq',1)}", styles["Normal"]))
        story.append(Spacer(1,8))

    # Short
    story.append(Paragraph("Section B - Short Answers", styles["Heading2"]))
    for q in data.get("shorts", []):
        story.append(Paragraph(f"Q{q['id']}: {q['question']}", styles["Normal"]))
        student_ans = student_answers.get("short", {}).get(str(q['id']), "")
        marks, sim = score_short_answer(student_ans, q.get("reference_answer",""))
        story.append(Paragraph(f"Reference Answer: {q.get('reference_answer','')}", styles["Normal"]))
        story.append(Paragraph(f"Student Answer: {student_ans}", styles["Normal"]))
        story.append(Paragraph(f"Marks: {marks:.2f}/{marks_config.get('short',5)} | Similarity: {sim*100:.1f}%", styles["Normal"]))
        story.append(Spacer(1,8))

    # Long
    story.append(Paragraph("Section C - Long Answers", styles["Heading2"]))
    for q in data.get("longs", []):
        story.append(Paragraph(f"Q{q['id']}: {q['question']}", styles["Normal"]))
        student_ans = student_answers.get("long", {}).get(str(q['id']), "")
        marks, sim = grade_long_answer(student_ans, q.get("reference_answer",""), marks_config.get("long",10))
        story.append(Paragraph(f"Reference Answer: {q.get('reference_answer','')}", styles["Normal"]))
        story.append(Paragraph(f"Student Answer: {student_ans}", styles["Normal"]))
        story.append(Paragraph(f"Marks: {marks:.2f}/{marks_config.get('long',10)} | Similarity: {sim*100:.1f}%", styles["Normal"]))
        story.append(Spacer(1,8))

    doc.build(story)
    buf.seek(0)
    return buf

# -------------------------
# DOCUMENT CREATION
# -------------------------
def make_docx(data, solved=False):
    doc = Document()
    doc.styles['Normal'].font.name = "Times New Roman"
    doc.styles['Normal'].font.size = Pt(11)
    doc.add_heading("AI Generated Question Paper", level=1)
    for q in data.get("mcqs", []):
        doc.add_paragraph(f"Q{q['id']}. {q.get('question','')}")
        for opt in q.get("options", []):
            doc.add_paragraph(f"- {opt}", style='List Bullet')
        if solved:
            letter = q.get("answer_letter","")
            try:
                correct = q.get("options", [])[ord(letter.lower())-97] if letter else ""
            except Exception:
                correct = ""
            doc.add_paragraph(f"✅ Answer: {correct}")
        doc.add_paragraph("")
    doc.add_heading("Short Questions", level=2)
    for q in data.get("shorts", []):
        doc.add_paragraph(f"Q{q['id']}. {q.get('question','')}")
        if solved:
            doc.add_paragraph(f"✅ Reference Answer: {q.get('reference_answer','')}")
        doc.add_paragraph("")
    doc.add_heading("Long Questions", level=2)
    for q in data.get("longs", []):
        doc.add_paragraph(f"Q{q['id']}. {q.get('question','')}")
        if solved:
            doc.add_paragraph(f"✅ Reference Answer: {q.get('reference_answer','')}")
        doc.add_paragraph("")
    bio = BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def make_pdf(data, solved=False):
    buf = BytesIO()
    doc = SimpleDocTemplate(buf)
    styles = getSampleStyleSheet()
    story = [Paragraph("AI Generated Question Paper", styles["Title"]), Spacer(1,12)]
    for q in data.get("mcqs", []):
        story.append(Paragraph(f"Q{q['id']}. {q.get('question','')}", styles["Normal"]))
        for opt in q.get("options", []):
            story.append(Paragraph(f"- {opt}", styles["Normal"]))
        if solved:
            letter = q.get("answer_letter","")
            try:
                correct = q.get("options", [])[ord(letter.lower())-97] if letter else ""
            except Exception:
                correct = ""
            story.append(Paragraph(f"<b>Answer:</b> {correct}", styles["Normal"]))
        story.append(Spacer(1,8))
    story.append(Paragraph("Short Questions", styles["Heading2"]))
    for q in data.get("shorts", []):
        story.append(Paragraph(f"Q{q['id']}. {q.get('question','')}", styles["Normal"]))
        if solved:
            story.append(Paragraph(f"<b>Reference Answer:</b> {q.get('reference_answer','')}", styles["Normal"]))
        story.append(Spacer(1,8))
    story.append(Paragraph("Long Questions", styles["Heading2"]))
    for q in data.get("longs", []):
        story.append(Paragraph(f"Q{q['id']}. {q.get('question','')}", styles["Normal"]))
        if solved:
            story.append(Paragraph(f"<b>Reference Answer:</b> {q.get('reference_answer','')}", styles["Normal"]))
        story.append(Spacer(1,8))
    doc.build(story)
    buf.seek(0)
    return buf
# -------------------------
# CSV + SESSION INITIALIZATION
# -------------------------
ROOMS_CSV = "rooms.csv"
PARTICIPANTS_CSV = "participants.csv"
QUIZ_RESULTS_CSV = "quiz_results.csv"

def ensure_csv(file, cols):
    if not os.path.exists(file):
        pd.DataFrame(columns=cols).to_csv(file, index=False)

ensure_csv(ROOMS_CSV, ["room_code","owner","created_at","status","paper_json","time_limit","start_time"])
ensure_csv(PARTICIPANTS_CSV, ["room_code","student_name","joined_at","avatar"])
ensure_csv(QUIZ_RESULTS_CSV, ["room_code","student_name","score","total","timestamp"])

def generate_room_code(length=6):
    return ''.join(random.choices(string.ascii_uppercase + string.digits, k=length))

AVATAR_EMOJIS = ["🐼","🦊","🐱","🐵","🐸","🦁","🐯","🐰","🐨","🐷","🐮","🐤"]

# -------------------------
# session state initialization (keys used)
# -------------------------
ss = st.session_state
_init_keys = {
    "uploaded_text": None,   # text extracted from uploaded PDF (Home)
    "uploaded_file_name": None,
    "paper": None,           # generated paper dict
    "answers": None,         # for Take Test per-user answers
    # Exam/Quiz specific
    "role": None,
    "active_room": None,
    "teacher_name": None,
    "joined_room": None,
    "joined_name": None,
    "joined_avatar": None,
    "quiz_qindex": 0,
    "quiz_score": 0,
    "quiz_in_progress": False,
    "last_join_event": None
}
for k,v in _init_keys.items():
    if k not in ss:
        ss[k] = v



def refresh_lobby(msg="Lobby refreshed!"):
    """
    Reusable function to refresh any lobby or room.
    """
    st.toast(f"🔁 {msg}")
    st.rerun()


# -------------------------
# RESULT DISPLAY HELPER
# -------------------------
def show_score_table(breakdown, total_score=None, total_marks=None):
    """Display result breakdown as a clean table"""
    rows = []
    for section, details in breakdown.items():
        for qid, marks in details.items():
            rows.append({
                "Section": section.capitalize(),
                "Question ID": qid,
                "Marks Obtained": marks
            })
    if rows:
        df_break = pd.DataFrame(rows)
        if total_score is not None and total_marks is not None:
            st.markdown(f"### 🎯 **Your Score:** {total_score} / {total_marks}")
        st.subheader("📊 Detailed Breakdown")
        st.table(df_break)
    else:
        st.info("No breakdown data available.")

# -------------------------
# SECTION SCORING HELPERS (reuses your existing functions)
# -------------------------

def score_mcq(student_answers, paper, marks_per_mcq=1):
    """Scores MCQs based on correct answer_letter"""
    total_score = 0
    breakdown = {}
    for q in paper.get("mcqs", []):
        qid = str(q["id"])
        student_choice = student_answers.get(qid, "").strip()
        letter = q.get("answer_letter", "").lower()
        correct = ""
        try:
            correct = q.get("options", [])[ord(letter) - 97] if letter else ""
        except Exception:
            pass
        if student_choice.lower() == correct.lower():
            total_score += marks_per_mcq
            breakdown[qid] = marks_per_mcq
        else:
            breakdown[qid] = 0
    return total_score, breakdown


def score_short_section(student_answers, paper, marks_per_short=5):
    """Uses your score_short_answer() for short-type questions"""
    total_score = 0
    breakdown = {}
    for q in paper.get("shorts", []):
        qid = str(q["id"])
        stu_ans = student_answers.get(qid, "")
        ref_ans = q.get("reference_answer", "")
        marks, sim = score_short_answer(stu_ans, ref_ans, max_marks=marks_per_short)
        total_score += marks
        breakdown[qid] = round(marks, 2)
    return total_score, breakdown


def score_long_section(student_answers, paper, marks_per_long=10):
    """Uses your grade_long_answer_fallback() for long-type questions"""
    total_score = 0
    breakdown = {}
    for q in paper.get("longs", []):
        qid = str(q["id"])
        stu_ans = student_answers.get(qid, "")
        ref_ans = q.get("reference_answer", "")
        marks, feedback = grade_long_answer_fallback(
            q.get("question", ""), ref_ans, stu_ans, max_marks=marks_per_long
        )
        total_score += marks
        breakdown[qid] = round(marks, 2)
    return total_score, breakdown


# -------------------------
# SAVE TO CSV (instead of SQLite)
# -------------------------
def save_result_csv(room_code, student_name, score, total, breakdown):
    try:
        df = pd.read_csv(QUIZ_RESULTS_CSV)
    except Exception:
        df = pd.DataFrame(columns=["room_code", "student_name", "score", "total", "breakdown", "timestamp"])

    df = pd.concat([
        df,
        pd.DataFrame([{
            "room_code": room_code,
            "student_name": student_name,
            "score": score,
            "total": total,
            "breakdown": json.dumps(breakdown),
            "timestamp": time.strftime("%Y-%m-%d %H:%M:%S")
        }])
    ], ignore_index=True)

    df.to_csv(QUIZ_RESULTS_CSV, index=False)

# -------------------------
# LOGIN PAGE
# -------------------------
if "logged_in" not in ss:
    ss.logged_in = False

if not ss.logged_in:
    st.markdown("<h1 style='text-align:center;'>🎓 EduGenius Hub Login</h1>", unsafe_allow_html=True)
    role = st.selectbox("Login as", ["Teacher", "Student"])
    password = st.text_input("Password", type="password")
    if role == "Teacher":
        st.caption("💡 Hint: Password is teach123")
    else:
        st.caption("💡 Hint: Password is stud123")
    if st.button("Login"):
        if (role == "Teacher" and password == TEACHER_PASS) or (role == "Student" and password == STUDENT_PASS):
            ss.logged_in = True
            ss.role = role
            st.success(f"Welcome, {role}!")
            st.rerun()
        else:
            st.error("Incorrect password.")


# ==============================
# PART 2: TEACHER & STUDENT DASHBOARDS
# ==============================

ss = st.session_state

if ss.logged_in:
    
    # -------------------------
    # TEACHER DASHBOARD
    # -------------------------
    if ss.role == "Teacher":
        st.title("👨‍🏫 Teacher Dashboard")
        menu = st.sidebar.selectbox("Menu", ["Generate Paper", "Conduct Exam", "Conduct Quick MCQs Quiz", "About"], key="teacher_menu")


        # Logout button below menu
        if st.sidebar.button("🔒 Logout"):
            # Clear session keys
            keys_to_clear = [
                "logged_in", "role", "teacher_name", "active_room", "paper",
                "answers", "joined_room", "joined_name", "joined_avatar",
                "quiz_qindex", "quiz_score", "quiz_in_progress", "last_join_event"
            ]
            for key in keys_to_clear:
                if key in ss:
                    del ss[key]

            st.success("Logged out successfully! Redirecting to login...")
            st.rerun()

         # -------------------------
        # Generate Paper
        # -------------------------
        if menu == "Generate Paper":
            st.subheader("1) Upload PDF / Notes")
            t_upload = st.file_uploader("Upload PDF (teacher) — used for generation", type=["pdf"], key="teacher_upload")
            if t_upload:
                try:
                    # Always extract text from newly uploaded PDF
                    txt = extract_text_from_pdf(t_upload)
                    ss.uploaded_text = txt
                    ss.uploaded_file_name = t_upload.name
                    st.success(f"PDF uploaded and text extracted: {t_upload.name}")
                except Exception as e:
                    st.error(f"Failed to extract text: {e}")

            st.subheader("2) Generate Question Paper")
            c1, c2, c3 = st.columns(3)
            with c1:
                n_mcq = st.number_input("Number of MCQs", min_value=0, max_value=50, value=5, key="t_n_mcq")
            with c2:
                n_short = st.number_input("Number of Short Qs", min_value=0, max_value=50, value=5, key="t_n_short")
            with c3:
                n_long = st.number_input("Number of Long Qs", min_value=0, max_value=20, value=2, key="t_n_long")

            m1, m2, m3 = st.columns(3)
            with m1:
                marks_mcq = st.number_input("Marks per MCQ", min_value=1, max_value=20, value=1, key="t_marks_mcq")
            with m2:
                marks_short = st.number_input("Marks per Short Q", min_value=1, max_value=50, value=5, key="t_marks_short")
            with m3:
                marks_long = st.number_input("Marks per Long Q", min_value=1, max_value=100, value=10, key="t_marks_long")

            style = st.text_input("Style / difficulty (optional)", key="t_style")

            if st.button("Generate Question Paper", key="t_generate"):
                source_text = ss.get("uploaded_text", "")
                if not source_text:
                    st.warning("No uploaded text in session — generator will use demo fallback.")
                with st.spinner("Generating..."):
                    try:
                        paper = generate_question_paper_from_text(source_text or "Sample fallback text",
                                                                n_mcq, n_short, n_long, style)
                        paper_meta = {"paper": paper,
                                    "marks": {"mcq": marks_mcq, "short": marks_short, "long": marks_long}}
                        ss.paper = paper_meta
                        st.success("Paper generated and stored in session.")
                    except Exception as e:
                        st.error(f"Generation failed: {e}")

            if ss.get("paper") and "paper" in ss.paper:
                st.subheader("Generated Paper Preview")
                pmeta = ss.paper
                p = pmeta["paper"]
                st.write(f"MCQs: {len(p.get('mcqs', []))} | Shorts: {len(p.get('shorts', []))} | Longs: {len(p.get('longs', []))}")

                with st.expander("Preview MCQs"):
                    for q in p.get("mcqs", []):
                        st.write(f"Q{q['id']}. {q.get('question')}")
                        for opt in q.get("options", []):
                            st.write(f"- {opt}")
                with st.expander("Preview Short Questions"):
                    for q in p.get("shorts", []):
                        st.write(f"Q{q['id']}. {q.get('question')}")
                with st.expander("Preview Long Questions"):
                    for q in p.get("longs", []):
                        st.write(f"Q{q['id']}. {q.get('question')}")

                q_doc = make_docx(p, solved=False)
                s_doc = make_docx(p, solved=True)
                q_pdf = make_pdf(p, solved=False)
                s_pdf = make_pdf(p, solved=True)

                colA, colB = st.columns(2)
                with colA:
                    st.download_button("⬇️ Question Paper (DOCX)", q_doc, file_name="Question_Paper.docx")
                    st.download_button("⬇️ Question Paper (PDF)", q_pdf, file_name="Question_Paper.pdf")
                with colB:
                    st.download_button("📘 Solved Paper (DOCX)", s_doc, file_name="Solved_Paper.docx")
                    st.download_button("📘 Solved Paper (PDF)", s_pdf, file_name="Solved_Paper.pdf")

        # -------------------------
        # Conduct Exam / Manage Rooms
        # -------------------------
        elif menu == "Conduct Exam":
            st.header("📝 Conduct Exam / Manage Rooms")

            # Teacher name input
            if "teacher_name" not in ss:
                ss.teacher_name = "Teacher"
            teacher_name = st.text_input("Teacher name", value=ss.teacher_name, key="teacher_name")

            # Room creation inputs
            room_code = st.text_input("Enter Room Code (or leave blank to generate new)")
            time_limit = st.number_input("Time Limit (minutes)", min_value=10, max_value=180, value=30)

            # Create Exam Room
            if st.button("Create Exam Room", key="t_create_room"):
                if not ss.get("paper"):
                    st.warning("Generate a paper first.")
                else:
                    code = room_code if room_code else generate_room_code()
                    try:
                        rooms = pd.read_csv(ROOMS_CSV)
                    except Exception:
                        rooms = pd.DataFrame(columns=["room_code","owner","created_at","status","paper_json","time_limit","start_time"])
                    
                    new_room = {
                        "room_code": code,
                        "owner": teacher_name,
                        "created_at": datetime.now(),
                        "status": "waiting",
                        "paper_json": json.dumps(ss["paper"]),
                        "time_limit": int(time_limit),
                        "start_time": ""
                    }
                    rooms = pd.concat([rooms, pd.DataFrame([new_room])], ignore_index=True)
                    rooms.to_csv(ROOMS_CSV, index=False)
                    ss.active_room = code
                    st.success(f"Exam Room created: {code} — share this code with students.")

            # --- Always reload rooms safely ---
            try:
                rooms = pd.read_csv(ROOMS_CSV)
            except Exception:
                rooms = pd.DataFrame(columns=["room_code","owner","created_at","status","paper_json","time_limit","start_time"])

            # Show all rooms
            if not rooms.empty:
                st.subheader("🏫 Active Rooms")
                st.dataframe(rooms[["room_code", "owner", "status", "created_at"]])
            else:
                st.info("No rooms available yet.")

            # --- Lobby / Active Room ---
            if ss.get("active_room"):
                st.markdown(f"**Active Room:** `{ss.active_room}`")
                st.subheader("🧑‍🎓 Students Lobby")

                # Rename button specifically for Conduct Exam Lobby
                conduct_exam = st.button("📋 Conduct Exam Lobby", key="conduct_exam_lobby")
                if conduct_exam:
                    st.toast("🔁 Conduct Exam Lobby refreshed!")
                    st.rerun()

                # refresh = st.button("🔄 Refresh Lobby", key="refresh_lobby")
                # if refresh:
                #     st.toast("🔁 Lobby refreshed!")
                #     st.rerun()

                # Load participants
                try:
                    parts = pd.read_csv(PARTICIPANTS_CSV)
                    joined = parts[parts["room_code"] == ss.active_room]
                    if not joined.empty:
                        st.success(f"{len(joined)} student(s) currently in the lobby.")
                        st.table(joined[["student_name", "avatar", "joined_at"]].rename(
                            columns={"student_name": "Name", "avatar": "Avatar", "joined_at": "Joined At"}
                        ))
                    else:
                        st.info("No students have joined yet.")
                except Exception:
                    st.warning("⚠️ Could not load participant data yet.")

                col1, col2 = st.columns(2)
                with col1:
                    if st.button("▶️ Start Exam", key="t_start_exam"):
                        try:
                            # Ensure rooms is reloaded fresh
                            rooms = pd.read_csv(ROOMS_CSV)
                            rooms.loc[rooms["room_code"] == ss.active_room, "status"] = "started"
                            rooms.loc[rooms["room_code"] == ss.active_room, "start_time"] = datetime.now()
                            rooms.to_csv(ROOMS_CSV, index=False)
                            st.success("✅ Exam started successfully!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to start exam: {e}")

                with col2:
                    if st.button("⛔ End Exam", key="t_end_exam"):
                        try:
                            rooms = pd.read_csv(ROOMS_CSV)
                            rooms.loc[rooms["room_code"] == ss.active_room, "status"] = "ended"
                            rooms.to_csv(ROOMS_CSV, index=False)
                            st.toast("🚫 Exam ended!")

                            # Store a flag in session to display results after buttons
                            ss["show_results"] = True
                        except Exception as e:
                            st.error(f"Failed to end exam: {e}")
                                # -------------------------
                # Move Results BELOW the buttons
                # -------------------------
                if ss.get("show_results"):
                    st.markdown("---")
                    st.header("📊 Final Scores & Breakdowns")

                    try:
                        res = pd.read_csv(QUIZ_RESULTS_CSV)
                        room_res = res[res["room_code"] == ss.active_room].sort_values(by="score", ascending=False)
                        if not room_res.empty:
                            for idx, row in room_res.iterrows():
                                st.markdown(f"### 👤 **{row['student_name']}** — {row['score']} / {row['total']}")
                                try:
                                    show_score_table(json.loads(row["breakdown"]))
                                except Exception:
                                    st.write("No breakdown available.")
                        else:
                            st.info("No submissions yet.")
                    except Exception:
                        st.info("No results yet.")

        # -------------------------
        # QUICK QUIZ — TEACHER SIDE (Enhanced with Timer + Marks)
        # -------------------------
        elif menu == "Conduct Quick MCQs Quiz":
            st.header("⚡ Quick Quiz — Teacher Panel")

            # Teacher info
            ss["teacher_name"] = st.text_input("Your name", value=ss.get("teacher_name") or "Teacher")

            c1, c2, c3 = st.columns(3)
            with c1:
                mcq_count = st.number_input("Number of MCQs", min_value=1, max_value=20, value=5)
            with c2:
                marks_per_mcq = st.number_input("Marks per MCQ", min_value=1, max_value=10, value=1)
            with c3:
                time_limit = st.number_input("⏱️ Quiz Duration (minutes)", min_value=1, max_value=60, value=5)

            start_quiz = st.button("🚀 Create Quick Quiz Room")

            if start_quiz:
                rc = generate_room_code(4)
                source_text = ss.get("uploaded_text", "Sample quick quiz text")

                # Generate paper
                try:
                    paper = generate_question_paper_from_text(source_text, num_mcq=mcq_count, num_short=0, num_long=0)
                except Exception:
                    paper = generate_question_paper_from_text("demo", num_mcq=mcq_count, num_short=0, num_long=0)

                # Fix JSON structure
                if "mcqs" in paper:
                    paper_json = json.dumps(paper)
                elif "paper" in paper and "mcqs" in paper["paper"]:
                    paper_json = json.dumps(paper["paper"])
                else:
                    paper_json = json.dumps({"mcqs": paper.get("mcqs", [])})

                # Save room
                try:
                    rooms = pd.read_csv(ROOMS_CSV)
                except Exception:
                    rooms = pd.DataFrame(columns=[
                        "room_code","owner","created_at","status","paper_json","time_limit","marks_per_mcq","start_time"
                    ])

                new_room = {
                    "room_code": rc,
                    "owner": ss.get("teacher_name", "Teacher"),
                    "created_at": datetime.now(),
                    "status": "waiting",
                    "paper_json": paper_json,
                    "time_limit": int(time_limit),
                    "marks_per_mcq": int(marks_per_mcq),
                    "start_time": ""
                }
                rooms = pd.concat([rooms, pd.DataFrame([new_room])], ignore_index=True)
                rooms.to_csv(ROOMS_CSV, index=False)
                ss["active_quick_room"] = rc
                st.success(f"✅ Quick Quiz Room created: `{rc}`")
                st.info(f"🕒 Time Limit: {time_limit} min | 🏅 Marks per MCQ: {marks_per_mcq}")

            # -------------------------
            # ACTIVE ROOM MANAGEMENT
            # -------------------------
            if ss.get("active_quick_room"):
                st.markdown(f"**Active Room:** `{ss['active_quick_room']}`")

                # === TIMER DISPLAY (LIVE) ===
                timer_placeholder = st.empty()
                try:
                    rooms = pd.read_csv(ROOMS_CSV)
                    room_row = rooms[rooms["room_code"] == ss["active_quick_room"]]
                    if not room_row.empty and str(room_row.iloc[-1]["start_time"]) not in ["", "nan"]:
                        start_time = pd.to_datetime(room_row.iloc[-1]["start_time"])
                        limit = int(room_row.iloc[-1].get("time_limit", 0))
                        if limit > 0:
                            end_time = start_time + pd.Timedelta(minutes=limit)
                            remaining = end_time - pd.Timestamp.now()
                            total_sec = remaining.total_seconds()
                            if total_sec > 0:
                                mins = int(total_sec // 60)
                                secs = int(total_sec % 60)
                                timer_placeholder.info(f"⏳ Time Remaining: {mins:02d}:{secs:02d}")
                            else:
                                timer_placeholder.warning("⏰ Quiz time over!")
                except Exception:
                    pass


                # === CONDUCT QUIZ LOBBY BUTTON ===
                # Auto-refresh every 3 seconds (3000 ms)
                #st_autorefresh(interval=3000, key="teacher_quiz_lobby_refresh")

                # Optional: show a toast/info to indicate refresh
                # st.info("📋 Quiz Lobby is auto-refreshing...")
                conduct_quiz = st.button("📋 Conduct Quiz Lobby", key="conduct_quiz_lobby")
                if conduct_quiz:
                    st.toast("Quiz Lobby refreshed!")
                    st.rerun()


                # # === REFRESH LOBBY BUTTON ===
                # refresh = st.button("🔄 Refresh Lobby")
                # if refresh:
                #     st.toast("Lobby refreshed!")
                #     st.rerun()

                # Show participants
                try:
                    parts = pd.read_csv(PARTICIPANTS_CSV)
                    joined = parts[parts["room_code"] == ss["active_quick_room"]]
                    if not joined.empty:
                        st.success(f"{len(joined)} student(s) currently in the lobby.")
                        st.table(joined[["student_name", "avatar", "joined_at"]].rename(
                            columns={"student_name": "Name", "avatar": "Avatar", "joined_at": "Joined At"}
                        ))
                    else:
                        st.info("No students have joined yet.")
                except Exception:
                    st.warning("⚠️ Could not load participant data yet.")



                # ---------------------------
                # Buttons for Start / End Quiz
                # ---------------------------
                col1, col2 = st.columns(2)
                with col1:
                    if st.button("▶️ Start Quick Quiz"):
                        try:
                            rooms = pd.read_csv(ROOMS_CSV)
                            rooms.loc[rooms["room_code"] == ss["active_quick_room"], "status"] = "started"
                            rooms.loc[rooms["room_code"] == ss["active_quick_room"], "start_time"] = datetime.now()
                            rooms.to_csv(ROOMS_CSV, index=False)
                            st.success("✅ Quick Quiz started!")
                            st.rerun()
                        except Exception as e:
                            st.error(f"Failed to start quiz: {e}")

                with col2:
                    end_quiz_clicked = st.button("⛔ End Room / Show Results")  # store click in variable

                # ---------------------------
                # Display Full-Width Results
                # ---------------------------
                if end_quiz_clicked:
                    try:
                        rooms = pd.read_csv(ROOMS_CSV)
                        rooms.loc[rooms["room_code"] == ss["active_quick_room"], "status"] = "ended"
                        rooms.to_csv(ROOMS_CSV, index=False)
                        st.toast("🚫 Quiz ended!")
                        st.markdown("---")

                        st.subheader("📊 Final Scores — Quick Quiz Results")
                        try:
                            res = pd.read_csv(QUIZ_RESULTS_CSV)
                            room_res = res[res["room_code"] == ss["active_quick_room"]].sort_values(by="score", ascending=False)
                            if not room_res.empty:
                                st.table(room_res[["student_name", "score", "total", "timestamp"]].rename(
                                    columns={"student_name": "Name", "score": "Score", "total": "Total Marks", "timestamp": "Submitted At"}
                                ))
                            else:
                                st.info("No student submissions yet.")
                        except Exception:
                            st.info("No quiz results available yet.")
                    except Exception as e:
                        st.error(f"Failed to end quiz: {e}")

        # -------------------------
        # ABOUT
        # -------------------------
        elif menu == "About":
            ### About the author
            st.write("##### About the author:")
                
            ### Author name
            st.write("<p style='color:blue; font-size: 50px; font-weight: bold;'>Usama Munawar</p>", unsafe_allow_html=True)
                
            ### Connect on social media
            st.write("##### Connect with me on social media")
                
            ### Add social media links
            ### URLs for images
            linkedin_url = "https://img.icons8.com/color/48/000000/linkedin.png"
            github_url = "https://img.icons8.com/fluent/48/000000/github.png"
            youtube_url = "https://img.icons8.com/?size=50&id=19318&format=png"
            twitter_url = "https://img.icons8.com/color/48/000000/twitter.png"
            facebook_url = "https://img.icons8.com/color/48/000000/facebook-new.png"
                
            ### Redirect URLs
            linkedin_redirect_url = "https://www.linkedin.com/in/abu--usama"
            github_redirect_url = "https://github.com/UsamaMunawarr"
            youtube_redirect_url ="https://www.youtube.com/@CodeBaseStats"
            twitter_redirect_url = "https://twitter.com/Usama__Munawar?t=Wk-zJ88ybkEhYJpWMbMheg&s=09"
            facebook_redirect_url = "https://www.facebook.com/profile.php?id=100005320726463&mibextid=9R9pXO"
                
            ### Add links to images
            st.markdown(f'<a href="{github_redirect_url}"><img src="{github_url}" width="60" height="60"></a>'
                        f'<a href="{linkedin_redirect_url}"><img src="{linkedin_url}" width="60" height="60"></a>'
                        f'<a href="{youtube_redirect_url}"><img src="{youtube_url}" width="60" height="60"></a>'
                        f'<a href="{twitter_redirect_url}"><img src="{twitter_url}" width="60" height="60"></a>'
                        f'<a href="{facebook_redirect_url}"><img src="{facebook_url}" width="60" height="60"></a>', unsafe_allow_html=True)
            # Thank you message
            st.write("<p style='color:green; font-size: 30px; font-weight: bold;'>Thank you for using this app, share with your friends!😇</p>", unsafe_allow_html=True)



    # -------------------------
    # STUDENT DASHBOARD
    # -------------------------
    elif ss.role == "Student":
        st.title("👨‍🎓 Student Dashboard")
        menu = st.sidebar.selectbox("Menu", ["Prepare Exam", "Attempt Exam", "Attempt Quick MCQs Quiz", "About"],key="student_menu")

        # Logout button below menu
        if st.sidebar.button("🔒 Logout"):
            # Clear session keys
            keys_to_clear = [
                "logged_in", "role", "teacher_name", "active_room", "paper",
                "answers", "joined_room", "joined_name", "joined_avatar",
                "quiz_qindex", "quiz_score", "quiz_in_progress", "last_join_event"
            ]
            for key in keys_to_clear:
                if key in ss:
                    del ss[key]

            st.success("Logged out successfully! Redirecting to login...")
            st.rerun()


        #---------------
        # Prepare Exam / Generate Paper
        #---------------
        if menu == "Prepare Exam":
            st.subheader("1) Upload PDF / Notes")
            s_upload = st.file_uploader("Upload PDF — used for generation", type=["pdf"], key="student_upload")
            if s_upload:
                try:
                    txt = extract_text_from_pdf(s_upload)
                    ss.uploaded_text = txt
                    ss.uploaded_file_name = s_upload.name
                    st.success("PDF uploaded and text extracted into session.")
                except Exception as e:
                    st.error(f"Failed to extract text: {e}")

            st.subheader("2) Find Important Questions")
            c1, c2, c3 = st.columns(3)
            with c1:
                n_mcq = st.number_input("Number of MCQs", min_value=0, max_value=50, value=5, key="s_n_mcq")
            with c2:
                n_short = st.number_input("Number of Short Qs", min_value=0, max_value=50, value=5, key="s_n_short")
            with c3:
                n_long = st.number_input("Number of Long Qs", min_value=0, max_value=20, value=2, key="s_n_long")

            style = st.text_input("Style / difficulty (optional)", key="s_style")

            if st.button("Generate Questions", key="s_generate"):
                source_text = ss.get("uploaded_text", "")
                if not source_text:
                    st.warning("No uploaded text in session — generator will use demo fallback.")
                with st.spinner("Generating..."):
                    try:
                        paper = generate_question_paper_from_text(
                            source_text or "Sample fallback text",
                            n_mcq, n_short, n_long, style
                        )
                        paper_meta = {"paper": paper}  # Marks removed for students
                        ss.paper = paper_meta
                        st.success("Paper generated and stored in session.")
                    except Exception as e:
                        st.error(f"Generation failed: {e}")

            if ss.get("paper") and "paper" in ss.paper:
                st.subheader("Generated Paper Preview")
                pmeta = ss.paper
                p = pmeta["paper"]
                st.write(f"MCQs: {len(p.get('mcqs', []))} | Shorts: {len(p.get('shorts', []))} | Longs: {len(p.get('longs', []))}")

                with st.expander("Preview MCQs"):
                    for q in p.get("mcqs", []):
                        st.write(f"Q{q['id']}. {q.get('question')}")
                        for opt in q.get("options", []):
                            st.write(f"- {opt}")
                with st.expander("Preview Short Questions"):
                    for q in p.get("shorts", []):
                        st.write(f"Q{q['id']}. {q.get('question')}")
                with st.expander("Preview Long Questions"):
                    for q in p.get("longs", []):
                        st.write(f"Q{q['id']}. {q.get('question')}")

                # Downloads
                q_doc = make_docx(p, solved=False)
                s_doc = make_docx(p, solved=True)
                q_pdf = make_pdf(p, solved=False)
                s_pdf = make_pdf(p, solved=True)

                colA, colB = st.columns(2)
                with colA:
                    st.download_button("⬇️ Question Paper (DOCX)", q_doc, file_name="Question_Paper.docx")
                    st.download_button("⬇️ Question Paper (PDF)", q_pdf, file_name="Question_Paper.pdf")
                with colB:
                    st.download_button("📘 Solved Paper (DOCX)", s_doc, file_name="Solved_Paper.docx")
                    st.download_button("📘 Solved Paper (PDF)", s_pdf, file_name="Solved_Paper.pdf")


        # -------------------------
        # Take Exam / Join Exam Room
        # -------------------------
        # -------------------------
        # Take Exam / Join or Generate Exam
        # -------------------------
        if menu == "Attempt Exam":
            st.subheader("Join Exam Room")
            name_in = st.text_input("Your name", key="join_name_student")
            code_in = st.text_input("Room code", key="join_code_student")

            if st.button("🔗 Join Room", key="student_join"):
                if not name_in or not code_in:
                    st.warning("Enter name and room code.")
                else:
                    try:
                        rooms = pd.read_csv(ROOMS_CSV)
                    except Exception:
                        rooms = pd.DataFrame(columns=["room_code","owner","created_at","status","paper_json","time_limit","start_time"])
                    matched = rooms[rooms["room_code"].astype(str) == str(code_in)]
                    if matched.empty:
                        st.error("Invalid or closed room code.")
                    else:
                        avatar = random.choice(AVATAR_EMOJIS)
                        try:
                            parts = pd.read_csv(PARTICIPANTS_CSV)
                        except Exception:
                            parts = pd.DataFrame(columns=["room_code","student_name","joined_at","avatar"])
                        new = {"room_code": code_in, "student_name": name_in, "joined_at": datetime.now(), "avatar": avatar}
                        parts = pd.concat([parts, pd.DataFrame([new])], ignore_index=True)
                        parts.to_csv(PARTICIPANTS_CSV, index=False)
                        ss["joined_room"] = code_in
                        ss["joined_name"] = name_in
                        ss["joined_avatar"] = avatar
                        st.success(f"Joined room {code_in} as {name_in} {avatar}")
                        st.rerun()

            # Exam Room / Lobby / Exam
            if ss.get("joined_room"):
                st.info(f"Joined Room `{ss['joined_room']}` as {ss['joined_name']} {ss['joined_avatar']}")

                # Load room info
                try:
                    rooms = pd.read_csv(ROOMS_CSV)
                    matched = rooms[rooms["room_code"] == ss["joined_room"]]
                    if not matched.empty:
                        status = matched.iloc[-1]["status"]
                        paper_json = matched.iloc[-1]["paper_json"]
                        time_limit = int(matched.iloc[-1].get("time_limit",30) or 30)
                        start_time_raw = matched.iloc[-1].get("start_time","")
                        start_time = pd.to_datetime(start_time_raw) if start_time_raw and str(start_time_raw)!="nan" else None
                    else:
                        status = "closed"
                        paper_json = None
                        start_time = None
                        time_limit = 0
                except Exception:
                    status = "closed"
                    paper_json = None
                    start_time = None
                    time_limit = 0

                if status == "waiting":
                    st.info("Waiting for the teacher to start the exam...")
                    if st.button("📝 Enter Exam Lobby", key="student_exam_lobby"):
                        st.rerun()

                elif status == "started":
                    # Countdown timer
                    if start_time:
                        end_time = start_time + pd.Timedelta(minutes=time_limit)
                        remaining = end_time - pd.Timestamp.now()
                        if remaining.total_seconds() <= 0:
                            status = "ended"
                            st.info("Time is up! Exam ended.")
                        else:
                            mins = int(remaining.total_seconds() // 60)
                            secs = int(remaining.total_seconds() % 60)
                            st.success(f"Exam in progress — time left: {mins}m {secs}s")
                    else:
                        st.success("Exam started.")

                    # Load paper
                    try:
                        papermeta = json.loads(paper_json) if paper_json else {"paper":{"mcqs":[],"shorts":[],"longs":[]}}
                    except Exception:
                        papermeta = {"paper":{"mcqs":[],"shorts":[],"longs":[]}}
                    paper = papermeta.get("paper",{})
                    marks_config = papermeta.get("marks", {"mcq":1,"short":5,"long":10})

                    # Initialize answers safely
                    # Initialize answers safely
                    if "answers" not in ss or ss["answers"] is None or not isinstance(ss["answers"], dict):
                        ss["answers"] = {"mcq": {}, "short": {}, "long": {}}
                    else:
                        ss["answers"].setdefault("mcq", {})
                        ss["answers"].setdefault("short", {})
                        ss["answers"].setdefault("long", {})

                    # MCQs
                    st.subheader("Section A - MCQs")
                    for q in paper.get("mcqs",[]):
                        key = f"exam_mcq_{q['id']}"
                        prev = ss["answers"]["mcq"].get(str(q['id']))
                        ss["answers"]["mcq"][str(q['id'])] = st.radio(
                            q.get("question",""),
                            options=q.get("options",[]),
                            index=0 if prev is None else q.get("options",[]).index(prev) if prev in q.get("options",[]) else 0,
                            key=key
                        )

                    # Short
                    st.subheader("Section B - Short Answers")
                    for q in paper.get("shorts",[]):
                        key = f"exam_short_{q['id']}"
                        ss["answers"]["short"][str(q['id'])] = st.text_area(
                            q.get("question",""), key=key, height=80,
                            value=ss["answers"]["short"].get(str(q['id']),"")
                        )

                    # Long
                    st.subheader("Section C - Long Answers")
                    for q in paper.get("longs",[]):
                        key = f"exam_long_{q['id']}"
                        ss["answers"]["long"][str(q['id'])] = st.text_area(
                            q.get("question",""), key=key, height=150,
                            value=ss["answers"]["long"].get(str(q['id']),"")
                        )

                    # Submit exam
                    if st.button("Submit Exam", key="student_submit_exam"):
                        with st.spinner("Grading your answers..."):
                            # Compute scores
                            mcq_score, mcq_break = score_mcq(ss["answers"]["mcq"], paper, marks_config.get("mcq",1))
                            short_score, short_break = score_short_section(ss["answers"]["short"], paper, marks_config.get("short",5))
                            long_score, long_break = score_long_section(ss["answers"]["long"], paper, marks_config.get("long",10))

                            total_score = mcq_score + short_score + long_score
                            total_marks = (marks_config.get("mcq",1)*len(paper.get("mcqs",[])) +
                                        marks_config.get("short",5)*len(paper.get("shorts",[])) +
                                        marks_config.get("long",10)*len(paper.get("longs",[])))

                            breakdown = {"mcq": mcq_break,"short": short_break,"long": long_break}

                            # Save results
                            try:
                                allr = pd.read_csv(QUIZ_RESULTS_CSV)
                            except Exception:
                                allr = pd.DataFrame(columns=["room_code","student_name","score","total","breakdown","timestamp"])

                            newr = {"room_code": ss["joined_room"], "student_name": ss["joined_name"],
                                    "score": total_score, "total": total_marks, "breakdown": json.dumps(breakdown),
                                    "timestamp": datetime.now()}
                            allr = pd.concat([allr, pd.DataFrame([newr])], ignore_index=True)
                            allr.to_csv(QUIZ_RESULTS_CSV, index=False)

                            st.success("✅ Exam submitted successfully!")
                            show_score_table(breakdown, total_score, total_marks)

                            # Reset answers
                            ss["answers"] = {"mcq":{}, "short":{}, "long":{}}




 
        # -------------------------
        # Quick MCQ Quiz Mode
        # -------------------------
        # # -------------------------
        # # STUDENT DASHBOARD
        # # -------------------------
        # -------------------------
        # STUDENT DASHBOARD
        # -------------------------
        if menu == "Attempt Quick MCQs Quiz":
            st.header("⚡ Quick MCQ Quiz — Student Panel")

            # If not joined, show join form
            if not ss.get("joined_quick_room"):
                c1, c2 = st.columns(2)
                with c1:
                    name_in = st.text_input("Your name", key="quick_name")
                with c2:
                    code_in = st.text_input("Room code", key="quick_code")

                if st.button("🔗 Join Quick Quiz"):
                    try:
                        rooms = pd.read_csv(ROOMS_CSV)
                    except Exception:
                        st.error("No active rooms found.")
                        rooms = pd.DataFrame(columns=["room_code", "owner", "created_at", "status", "paper_json"])

                    matched = rooms[rooms["room_code"].astype(str) == str(code_in)]
                    if matched.empty:
                        st.error("Invalid or closed room code.")
                    else:
                        avatar = random.choice(AVATAR_EMOJIS)
                        try:
                            parts = pd.read_csv(PARTICIPANTS_CSV)
                        except Exception:
                            parts = pd.DataFrame(columns=["room_code","student_name","joined_at","avatar"])

                        new = {"room_code": code_in, "student_name": name_in, "joined_at": datetime.now(), "avatar": avatar}
                        parts = pd.concat([parts, pd.DataFrame([new])], ignore_index=True)
                        parts.to_csv(PARTICIPANTS_CSV, index=False)

                        ss["joined_quick_room"] = code_in
                        ss["joined_name_quick"] = name_in
                        ss["joined_avatar_quick"] = avatar
                        ss["quiz_qindex"] = 0
                        ss["quiz_score"] = 0
                        st.success(f"Joined room {code_in} as {name_in} {avatar}")
                        time.sleep(0.5)
                        st.rerun()

            # If already joined
            else:
                st.info(f"Joined Quick Quiz Room `{ss['joined_quick_room']}` as {ss['joined_name_quick']} {ss['joined_avatar_quick']}")

                # Exit button
                if st.button("🚪 Leave Lobby / Exit Room", key="leave_quick_lobby"):
                    ss["joined_quick_room"] = None
                    ss["joined_name_quick"] = None
                    ss["joined_avatar_quick"] = None
                    ss["quiz_qindex"] = 0
                    ss["quiz_score"] = 0
                    st.success("You left the lobby.")
                    st.rerun()

                # Check if quiz started
                try:
                    rooms = pd.read_csv(ROOMS_CSV)
                    matched = rooms[rooms["room_code"] == ss["joined_quick_room"]]
                    if not matched.empty and matched.iloc[-1]["status"] == "started":
                        paper_json = matched.iloc[-1]["paper_json"]
                        started = True
                    else:
                        started = False
                except Exception:
                    started = False

                # Waiting lobby
                if not started:
                    st.info("⏳ Waiting for teacher to start the Quick Quiz...")
                    #if st.button("🔄 Refresh Lobby"):
                        #st.rerun()
                    # 🕒 Auto-refresh every 3 seconds until quiz starts
                    st_autorefresh(interval=3000, key="student_quickquiz_lobby_autorefresh")
                    st.caption("🎯 Lobby auto-refreshes every 3 seconds — quiz will start automatically once the teacher begins.")
                # Quiz started
                else:
                    st.success("✅ Quiz started!")

                    # === TIMER (LIVE + AUTO-SUBMIT) ===
                    timer_placeholder = st.empty()
                    try:
                        start_time = pd.to_datetime(matched.iloc[-1]["start_time"])
                        limit = int(matched.iloc[-1].get("time_limit", 0))
                        if limit > 0:
                            end_time = start_time + pd.Timedelta(minutes=limit)
                            remaining = end_time - pd.Timestamp.now()
                            total_sec = remaining.total_seconds()

                            if total_sec <= 0:
                                st.error("⏰ Time is up! Auto-submitting your quiz…")
                                total = len(json.loads(paper_json).get("mcqs", []))
                                score = ss.get("quiz_score", 0)
                                try:
                                    allr = pd.read_csv(QUIZ_RESULTS_CSV)
                                except Exception:
                                    allr = pd.DataFrame(columns=["room_code","student_name","score","total","timestamp"])
                                new_row = {
                                    "room_code": ss["joined_quick_room"],
                                    "student_name": ss["joined_name_quick"],
                                    "score": score,
                                    "total": total,
                                    "timestamp": datetime.now()
                                }
                                allr = pd.concat([allr, pd.DataFrame([new_row])], ignore_index=True)
                                allr.to_csv(QUIZ_RESULTS_CSV, index=False)
                                st.success("✅ Quiz auto-submitted!")
                                ss["joined_quick_room"] = None
                                ss["joined_name_quick"] = None
                                ss["joined_avatar_quick"] = None
                                ss["quiz_qindex"] = 0
                                ss["quiz_score"] = 0
                                st.rerun()
                            else:
                                mins = int(total_sec // 60)
                                secs = int(total_sec % 60)
                                timer_placeholder.info(f"⏳ Time Remaining: {mins:02d}:{secs:02d}")
                                time.sleep(1)
                                st.experimental_rerun()
                    except Exception:
                        pass

                    # Load paper
                    try:
                        paper = json.loads(paper_json)
                    except Exception:
                        paper = {"mcqs":[]}
                    mcqs = paper.get("mcqs", [])
                    total = len(mcqs)
                    qidx = ss.get("quiz_qindex", 0)

                    if qidx >= total:
                        # ✅ Auto-finish immediately after last question
                        st.success("🎉 Quiz Finished & Submitted Automatically!")

                        # Compute total score and save result
                        total = len(mcqs)
                        score = ss.get("quiz_score", 0)

                        # Save result to CSV
                        try:
                            allr = pd.read_csv(QUIZ_RESULTS_CSV)
                        except Exception:
                            allr = pd.DataFrame(columns=["room_code","student_name","score","total","timestamp"])

                        exists = (
                            (allr["room_code"].astype(str) == str(ss["joined_quick_room"])) &
                            (allr["student_name"].astype(str) == str(ss["joined_name_quick"]))
                        ).any() if not allr.empty else False

                        if not exists:
                            new_row = {
                                "room_code": ss["joined_quick_room"],
                                "student_name": ss["joined_name_quick"],
                                "score": score,
                                "total": total,
                                "timestamp": datetime.now()
                            }
                            allr = pd.concat([allr, pd.DataFrame([new_row])], ignore_index=True)
                            allr.to_csv(QUIZ_RESULTS_CSV, index=False)

                        # 🧹 Clear lobby session (exit automatically)
                        current_room = ss["joined_quick_room"]
                        ss["joined_quick_room"] = None
                        ss["joined_name_quick"] = None
                        ss["joined_avatar_quick"] = None
                        ss["quiz_qindex"] = 0
                        ss["quiz_score"] = 0

                        # 🏁 Show Final Score and Leaderboard Instantly
                        st.markdown("---")
                        st.header("📊 Your Final Results")

                        st.success(f"**Your Score:** {score} / {total}")

                        # Try to show leaderboard for that room
                        try:
                            res = pd.read_csv(QUIZ_RESULTS_CSV)
                            room_res = res[res["room_code"] == current_room].sort_values(by="score", ascending=False)
                            if not room_res.empty:
                                st.subheader("🏆 Leaderboard")
                                st.table(room_res[["student_name","score","total","timestamp"]].rename(
                                    columns={"student_name": "Name", "score": "Score", "total": "Total Marks", "timestamp": "Submitted At"}
                                ))
                            else:
                                st.info("No other submissions yet.")
                        except Exception:
                            st.info("No leaderboard data available yet.")

                        st.info("✅ You can now return to Home or wait for the teacher to view all results.")

                    else:
                        q = mcqs[qidx]
                        st.write(f"**Q{q.get('id','')}**: {q.get('question','')}")
                        opts = q.get("options", [])
                        key_choice = f"quick_choice_{ss['joined_quick_room']}_{ss['joined_name_quick']}_{q.get('id')}"
                        choice = st.radio("Choose answer", options=opts, key=key_choice)
                        if st.button("Submit Answer"):
                            letter = q.get("answer_letter","")
                            correct_text = opts[ord(letter.lower())-97] if letter else q.get("answer","")
                            correct_flag = (choice.strip().lower() == correct_text.strip().lower())
                            ss["quiz_qindex"] += 1
                            time.sleep(0.5)
                            st.rerun()


        # -------------------------
        # ABOUT
        # -------------------------

        elif menu == "About":
            ### About the author
            st.write("##### About the author:")
                
            ### Author name
            st.write("<p style='color:blue; font-size: 50px; font-weight: bold;'>Usama Munawar</p>", unsafe_allow_html=True)
                
            ### Connect on social media
            st.write("##### Connect with me on social media")
                
            ### Add social media links
            ### URLs for images
            linkedin_url = "https://img.icons8.com/color/48/000000/linkedin.png"
            github_url = "https://img.icons8.com/fluent/48/000000/github.png"
            youtube_url = "https://img.icons8.com/?size=50&id=19318&format=png"
            twitter_url = "https://img.icons8.com/color/48/000000/twitter.png"
            facebook_url = "https://img.icons8.com/color/48/000000/facebook-new.png"
                
            ### Redirect URLs
            linkedin_redirect_url = "https://www.linkedin.com/in/abu--usama"
            github_redirect_url = "https://github.com/UsamaMunawarr"
            youtube_redirect_url ="https://www.youtube.com/@CodeBaseStats"
            twitter_redirect_url = "https://twitter.com/Usama__Munawar?t=Wk-zJ88ybkEhYJpWMbMheg&s=09"
            facebook_redirect_url = "https://www.facebook.com/profile.php?id=100005320726463&mibextid=9R9pXO"
                
            ### Add links to images
            st.markdown(f'<a href="{github_redirect_url}"><img src="{github_url}" width="60" height="60"></a>'
                        f'<a href="{linkedin_redirect_url}"><img src="{linkedin_url}" width="60" height="60"></a>'
                        f'<a href="{youtube_redirect_url}"><img src="{youtube_url}" width="60" height="60"></a>'
                        f'<a href="{twitter_redirect_url}"><img src="{twitter_url}" width="60" height="60"></a>'
                        f'<a href="{facebook_redirect_url}"><img src="{facebook_url}" width="60" height="60"></a>', unsafe_allow_html=True)
            # Thank you message
            st.write("<p style='color:green; font-size: 30px; font-weight: bold;'>Thank you for using this app, share with your friends!😇</p>", unsafe_allow_html=True)

